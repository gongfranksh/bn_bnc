#coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import os, sys


class bnc_report_word(object):

    def create_word(self, data):

        # 打开文档

        title=data['title']
        categorys=data['files']
        volumns=data['volumn']


        document = Document()
        # 加入不同等级的标题
        document.add_heading(title, 0)
        for cat in categorys:
            document.add_heading(cat['name'])
            file_path = os.path.join(sys.path[0], cat['filename'])
            document.add_picture(file_path, width=Inches(7.5))

        # document.add_heading(u'一级标题', 1)
        # document.add_heading(u'二级标题', 2)
        # 添加文本
        paragraph = document.add_paragraph(u'我们在做文本测试！')
        # 设置字号
        run = paragraph.add_run(u'设置字号、')
        run.font.size = Pt(24)

        # 设置字体
        run = paragraph.add_run('Set Font,')
        run.font.name = 'Consolas'

        # 设置中文字体
        run = paragraph.add_run(u'设置中文字体、')
        run.font.name = u'宋体'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # 设置斜体
        run = paragraph.add_run(u'斜体、')
        run.italic = True

        # 设置粗体
        run = paragraph.add_run(u'粗体').bold = True

        # 增加引用
        document.add_paragraph('Intense quote', style='Intense Quote')

        # 增加无序列表
        document.add_paragraph(
            u'无序列表元素1', style='List Bullet'
        )
        document.add_paragraph(
            u'无序列表元素2', style='List Bullet'
        )
        # 增加有序列表
        document.add_paragraph(
            u'有序列表元素1', style='List Number'
        )
        document.add_paragraph(
            u'有序列表元素2', style='List Number'
        )
        # 增加图像（此处用到图像image.bmp，请自行添加脚本所在目录中）
        #filename = u'2018-001-2018年度1至4月-0.png'
        # filename = u'%s-%s-%s.png' % (self.code, self.name, i)

        # 增加表格
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        # 再增加3行表格元素
        for i in xrange(3):
            row_cells = table.add_row().cells
            row_cells[0].text = 'test' + str(i)
            row_cells[1].text = str(i)
            row_cells[2].text = 'desc' + str(i)

        # 增加分页
        document.add_page_break()

        # 保存文件
        document.save(u'测试.docx')

        filename = u'测试.docx'
        # filename = u'%s-%s-%s.png' % (self.code, self.name, i)
        file_path = os.path.join(sys.path[0], filename)

        with open(file_path, 'rb') as fp:
            data = fp.read().encode('base64')


        return data
