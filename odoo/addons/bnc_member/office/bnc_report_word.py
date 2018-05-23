# coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import os, sys


class bnc_report_word(object):

    def create_word(self, data):

        # 打开文档

        title = data['title']
        categorys = data['files']

        document = Document()
        # 加入不同等级的标题
        document.add_heading(title, 0)
        for cat in categorys:
            document.add_heading(cat['name'])
            file_path = os.path.join(sys.path[0], cat['filename'])
            document.add_picture(file_path, width=Inches(7.5))

        document.styles['Normal'].font.name = 'Consolas'
        document.styles['Normal'].font.size = Pt(8)
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'名称'
        hdr_cells[1].text = u'聚类'
        hdr_cells[2].text = u'会员数量'
        # 再增加3行表格元素
        for cat in categorys:
            row_cells = table.add_row().cells
            row_cells[0].text = cat['name']
            row_cells[1].text = str(cat['category'])
            row_cells[2].text = str(cat['volumns'])


        # 增加分页
        document.add_page_break()

        # 保存文件
        filename = title +'.docx'
        file_path = os.path.join(sys.path[0], filename)

        if os.path.exists(file_path):
            os.remove(file_path)

        document.save(file_path)

        with open(file_path, 'rb') as fp:
            data = fp.read().encode('base64')

        return data
